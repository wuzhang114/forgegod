# -*- coding: utf-8 -*-
"""从 data/mechanism-library.json 生成《锻造神赐机制蓝图库 v1》docx 文档。"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
JSON_PATH = os.path.join(BASE, "data", "mechanism-library.json")
OUT_PATH = os.path.join(BASE, "锻造神赐机制蓝图库-v1.docx")

FONT = "微软雅黑"
TIER_TEXT = {"A": "Tier A · 首版原语可实现", "B": "Tier B · 需扩展原语", "C": "Tier C · 小队/世界级"}


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)


def set_style(doc, name, size, bold=False, color=None, all_caps=False):
    st = doc.styles[name]
    st.font.name = FONT
    st.font.size = Pt(size)
    st.font.bold = bold
    if color:
        st.font.color.rgb = color
    if all_caps:
        st.font.small_caps = True
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)


def para(doc, text, size=10.5, bold=False, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(doc, text, level=1, color=None):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_run(r, size=16 if level == 1 else 13 if level == 2 else 11.5, bold=True, color=color)
    return h


def kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.autofit = False
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(2.6)
        c1.width = Cm(14.6)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        set_run(r0, size=10, bold=True)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        set_run(r1, size=10)
        p0.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
    return t


def main():
    with open(JSON_PATH, encoding="utf-8") as f:
        data = json.load(f)
    meta = data["meta"]
    doc = Document()
    for name, size, bold, color in [
        ("Normal", 10.5, False, None),
        ("Heading 1", 16, True, RGBColor(0x8B, 0x2F, 0x1A)),
        ("Heading 2", 13, True, RGBColor(0x4A, 0x33, 0x22)),
        ("Heading 3", 11.5, True, RGBColor(0x1F, 0x4E, 0x79)),
        ("Title", 24, True, RGBColor(0x8B, 0x2F, 0x1A)),
    ]:
        set_style(doc, name, size, bold, color)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)

    # 封面区
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(meta["title"])
    set_run(r, size=26, bold=True, color=RGBColor(0x8B, 0x2F, 0x1A))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta["subtitle"])
    set_run(r, size=12, color=RGBColor(0x55, 0x55, 0x55))
    para(doc, "版本 v1 · %s" % meta["date"], size=10, color=RGBColor(0x88, 0x88, 0x88),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 说明
    heading(doc, "说明与用途", 1)
    para(doc, meta["intent"], size=10.5)
    para(doc, "分级说明:", size=10.5, bold=True)
    for tier in meta["tiers"]:
        para(doc, "· %s — %s" % (TIER_TEXT[tier["id"]], tier["meaning"]), size=10)

    # 分类条目
    for cat in data["categories"]:
        heading(doc, cat["title"], 1)
        para(doc, cat["intro"], size=10.5, color=RGBColor(0x44, 0x44, 0x44))
        for item in cat["items"]:
            heading(doc, "%s · %s（%s）" % (item["id"], item["name"], TIER_TEXT[item["tier"]]), 3)
            rows = [
                ("一句话幻想", item["fantasy"]),
                ("玩家申请", item["application"]),
                ("机制构成", item["mechanism"]),
                ("器物依据", item["basis"]),
                ("神前质询", item["question"]),
                ("代价方案", item["costs"]),
                ("反制入口", item["counter"]),
                ("战报标签", item["report_tags"]),
                ("原语组合", item["ops"]),
            ]
            kv_table(doc, rows)
            para(doc, "", size=6)

    # 附录
    heading(doc, "附录一：使用说明", 1)
    for note in data["appendix"]["usage_notes"]:
        para(doc, note, size=10)
    heading(doc, "附录二：按实施优先级分组", 1)
    for tier_id in ["A", "B", "C"]:
        items = data["appendix"]["priority"][tier_id + "_items"]
        para(doc, "%s: %s" % (TIER_TEXT[tier_id], "、".join(items)), size=10)

    doc.save(OUT_PATH)
    print("SAVED:", OUT_PATH)
    n = sum(len(c["items"]) for c in data["categories"])
    print("TOTAL_ITEMS:", n)


if __name__ == "__main__":
    main()
